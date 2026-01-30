"""
文档处理引擎 - 提供 Word 文档创建和编辑功能
"""
from typing import Optional, List, Dict, Any
from pathlib import Path
from everify.core.utils import logger
from everify.core.utils import config


class DocumentEngine:
    """文档处理引擎接口"""

    def __init__(self, document_config: Optional[Any] = None):
        self.config = document_config or config
        logger.debug("文档处理引擎初始化")

    def create_document(self, title: str, output_path: Optional[Path] = None) -> Path:
        """创建新文档"""
        pass

    def add_title(self, title: str, level: int = 1) -> None:
        """添加标题"""
        pass

    def add_paragraph(self, text: str) -> None:
        """添加段落"""
        pass

    def add_image(self, image_path: Path, caption: Optional[str] = None, width: Optional[int] = None) -> None:
        """添加图片"""
        pass

    def add_table(self, data: List[List[Any]], headers: Optional[List[str]] = None) -> None:
        """添加表格"""
        pass

    def add_section(self, title: str) -> None:
        """添加章节"""
        pass

    def save_document(self, output_path: Optional[Path] = None) -> Path:
        """保存文档"""
        pass

    def close_document(self) -> None:
        """关闭文档"""
        pass


class DocxDocumentEngine(DocumentEngine):
    """基于 python-docx 的文档处理引擎实现"""

    def __init__(self, document_config: Optional[Any] = None):
        super().__init__(document_config)
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.oxml.ns import qn
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.shared import RGBColor

        self.Document = Document
        self.Inches = Inches
        self.Pt = Pt
        self.qn = qn
        self.WD_PARAGRAPH_ALIGNMENT = WD_PARAGRAPH_ALIGNMENT
        self.RGBColor = RGBColor

        self.doc = None
        self.current_output_path = None
        self.level2_counter = 0
        logger.debug("DocxDocumentEngine 初始化完成")

    def create_document(self, title: str, output_path: Optional[Path] = None) -> Path:
        """创建新文档"""
        self.doc = self.Document()
        self.level2_counter = 0  # 重置二级标题编号

        # 设置文档默认属性
        self.doc.core_properties.title = title
        self.doc.core_properties.author = "Everify"
        self.doc.core_properties.subject = "诚信核查报告"

        if output_path is None:
            from datetime import datetime
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = self.config.reports_dir / f"{timestamp}_report.docx"

        self.current_output_path = output_path
        logger.debug(f"创建新文档: {output_path}")

        return output_path

    def add_title(self, title: str, level: int = 1) -> None:
        """添加标题"""
        if not self.doc:
            logger.warning("文档尚未创建，无法添加标题")
            return

        try:
            # 为二级标题添加编号
            if level == 2:
                self.level2_counter += 1
                title = f"{self.level2_counter}、{title}"

            heading = self.doc.add_heading(title, level=level)

            if level == 1:
                # 一级标题格式：四号字体，居中对齐，1.5倍行距，首行不缩进，黑色
                heading.paragraph_format.alignment = self.WD_PARAGRAPH_ALIGNMENT.CENTER
                heading.paragraph_format.line_spacing = 1.5
                heading.paragraph_format.first_line_indent = 0

                for run in heading.runs:
                    run.font.name = '楷体'
                    run._element.rPr.rFonts.set(self.qn('w:eastAsia'), '楷体')
                    run.font.size = self.Pt(14)  # 四号字体
                    run.font.color.rgb = self.RGBColor(0, 0, 0)  # 黑色
            elif level == 2:
                # 二级标题格式：小四字体，左对齐，1.5倍行距，首行缩进，黑色，标题后不换行
                heading.paragraph_format.alignment = self.WD_PARAGRAPH_ALIGNMENT.LEFT
                heading.paragraph_format.line_spacing = 1.5
                heading.paragraph_format.first_line_indent = self.Pt(14)  # 首行缩进2字符（小四字体）
                heading.paragraph_format.space_after = self.Pt(0)  # 标题后不换行

                for run in heading.runs:
                    if any('\u4e00' <= c <= '\u9fff' for c in run.text):  # 包含汉字
                        run.font.name = '楷体'
                        run._element.rPr.rFonts.set(self.qn('w:eastAsia'), '楷体')
                    else:  # 英文或数字
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(self.qn('w:ascii'), 'Times New Roman')
                        run._element.rPr.rFonts.set(self.qn('w:hAnsi'), 'Times New Roman')
                    run.font.size = self.Pt(12)  # 小四字体
                    run.font.color.rgb = self.RGBColor(0, 0, 0)  # 黑色

            logger.debug(f"添加标题: {title} (级别: {level})")
        except Exception as e:
            logger.error(f"添加标题失败: {e}")

    def add_paragraph(self, text: str) -> None:
        """添加段落"""
        if not self.doc:
            logger.warning("文档尚未创建，无法添加段落")
            return

        try:
            para = self.doc.add_paragraph(text)

            # 设置段落格式：小四字体，1.5倍行距，首行缩进
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.first_line_indent = self.Pt(14)  # 首行缩进2字符（小四字体）

            # 设置字体
            for run in para.runs:
                if any('\u4e00' <= c <= '\u9fff' for c in run.text):  # 包含汉字
                    run.font.name = '楷体'
                    run._element.rPr.rFonts.set(self.qn('w:eastAsia'), '楷体')
                else:  # 英文或数字
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(self.qn('w:ascii'), 'Times New Roman')
                    run._element.rPr.rFonts.set(self.qn('w:hAnsi'), 'Times New Roman')
                run.font.size = self.Pt(12)  # 小四字体

            para.paragraph_format.space_after = self.Pt(6)
            logger.debug(f"添加段落: {text[:50]}...")
        except Exception as e:
            logger.error(f"添加段落失败: {e}")

    def add_image(self, image_path: Path, caption: Optional[str] = None, width: Optional[int] = None) -> None:
        """添加图片"""
        if not self.doc:
            logger.warning("文档尚未创建，无法添加图片")
            return

        if not image_path.exists():
            logger.warning(f"图片文件不存在: {image_path}")
            self.add_paragraph("图片未找到")
            return

        try:
            # 确定图片宽度
            img_width = width if width is not None else self.Inches(6)

            # 添加图片
            para = self.doc.paragraphs[-1]  # 获取当前段落
            run = para.add_run()
            run.add_picture(str(image_path), width=img_width)

            # 设置图片居中对齐
            para.paragraph_format.alignment = self.WD_PARAGRAPH_ALIGNMENT.CENTER

            # 添加图片说明
            if caption:
                self.doc.add_paragraph(caption, style='Caption')

            logger.debug(f"添加图片: {image_path}")
        except Exception as e:
            logger.error(f"添加图片失败: {e}")
            self.add_paragraph(f"图片加载失败: {image_path}")

    def add_table(self, data: List[List[Any]], headers: Optional[List[str]] = None) -> None:
        """添加表格"""
        if not self.doc:
            logger.warning("文档尚未创建，无法添加表格")
            return

        try:
            # 计算列数
            if headers:
                col_count = len(headers)
            elif data:
                col_count = len(data[0])
            else:
                logger.warning("表格数据为空")
                return

            # 创建表格
            if headers:
                table = self.doc.add_table(rows=1, cols=col_count)
                hdr_cells = table.rows[0].cells

                # 添加表头
                for i, header in enumerate(headers):
                    hdr_cells[i].text = str(header)

                # 设置表头样式
                for cell in hdr_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
            else:
                table = self.doc.add_table(rows=0, cols=col_count)

            # 添加数据行
            for row_data in data:
                row_cells = table.add_row().cells
                for i, cell_data in enumerate(row_data):
                    if i < len(row_cells):
                        row_cells[i].text = str(cell_data)

            # 设置表格样式
            table.style = 'Table Grid'

            logger.debug(f"添加表格: {len(data)} 行, {col_count} 列")
        except Exception as e:
            logger.error(f"添加表格失败: {e}")

    def add_section(self, title: str) -> None:
        """添加章节"""
        if not self.doc:
            logger.warning("文档尚未创建，无法添加章节")
            return

        try:
            # 添加章节分隔符
            self.doc.add_page_break()
            self.add_title(title, level=1)
            logger.debug(f"添加章节: {title}")
        except Exception as e:
            logger.error(f"添加章节失败: {e}")

    def save_document(self, output_path: Optional[Path] = None) -> Path:
        """保存文档"""
        if not self.doc:
            logger.warning("文档尚未创建，无法保存")
            return None

        try:
            if output_path is None:
                if self.current_output_path is None:
                    from datetime import datetime
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    output_path = self.config.reports_dir / f"{timestamp}_report.docx"
                else:
                    output_path = self.current_output_path
            else:
                # 直接使用传入的路径，不进行额外拼接，避免 reports/reports 嵌套
                output_path = Path(output_path)

            output_path.parent.mkdir(exist_ok=True, parents=True)
            self.doc.save(str(output_path))
            logger.info(f"文档保存成功: {output_path}")

            return output_path
        except Exception as e:
            logger.error(f"保存文档失败: {e}")
            return None

    def close_document(self) -> None:
        """关闭文档"""
        if self.doc:
            try:
                self.doc = None
                self.current_output_path = None
                logger.debug("文档已关闭")
            except Exception as e:
                logger.error(f"关闭文档失败: {e}")


def create_document_engine() -> DocumentEngine:
    """创建文档处理引擎实例"""
    try:
        return DocxDocumentEngine()
    except Exception as e:
        logger.error(f"创建文档引擎失败: {e}")
        raise
