#!/usr/bin/env python3
"""
报告生成服务模块
负责生成Word报告
"""
from typing import Dict, List, Optional
from pathlib import Path
from everify.core.utils import logger
from everify.core.base.document import DocumentEngine, DocxDocumentEngine
from everify.core.base.image import ImageEngine, PillowImageEngine
from everify.core.utils.config import VerifyTemplate, WatermarkConfig


class ReportGenerator:
    """报告生成服务"""

    def __init__(self, config=None):
        """初始化报告生成器

        Args:
            config: 配置对象
        """
        self.config = config
        self.document_engine: DocumentEngine = DocxDocumentEngine()
        self.image_engine: ImageEngine = PillowImageEngine()
        self.templates: Optional[Dict[str, VerifyTemplate]] = None

    def generate_report(
        self,
        results: Dict[str, Dict[str, str]],
        output_path: Optional[Path] = None,
        templates: Optional[Dict[str, VerifyTemplate]] = None
    ) -> Dict[str, Path]:
        """生成最终的 Word 报告

        Args:
            results: 核查结果 {主体名称: {URL: 截图路径}}
            output_path: 报告输出路径
            templates: 核查模板字典

        Returns:
            dict: {主体名称: 报告文件路径}
        """
        if templates:
            self.templates = templates

        output_dir = output_path or self.config.reports_dir
        output_dir.mkdir(parents=True, exist_ok=True)

        report_paths = {}

        for entity, entity_results in results.items():
            try:
                report_path = self._generate_single_report(entity, entity_results, output_dir)
                report_paths[entity] = report_path
                logger.info(f"主体 '{entity}' 报告生成成功: {report_path}")
            except Exception as e:
                logger.error(f"主体 '{entity}' 报告生成失败: {e}")

        return report_paths

    def _generate_single_report(
        self,
        entity: str,
        entity_results: Dict[str, str],
        output_dir: Path
    ) -> Path:
        """为单个主体生成报告

        Args:
            entity: 主体名称
            entity_results: 主体的核查结果 {URL: 截图路径}
            output_dir: 输出目录

        Returns:
            Path: 报告文件路径
        """
        from everify.utils.file import clean_filename

        # 创建文档
        from datetime import datetime
        current_date = datetime.now().strftime("%Y%m%d")
        title = f"{entity}诚信核查报告{current_date}"
        self.document_engine.create_document(title)

        # 添加标题
        self.document_engine.add_title(title, level=1)


        # 定义可以自动化核查的网页
        automated_chapters = [
            "百度（https://www.baidu.com/s）",
            "中华人民共和国生态环境部网站（https://www.mee.gov.cn）",
            "中华人民共和国商务部网站（https://search.mofcom.gov.cn）",
            "国家外汇管理局网站（http://www.safe.gov.cn）",
            "中国人民银行网站（https://wzdig.pbc.gov.cn）",
            "国家金融监督管理总局（https://www.nfra.gov.cn/cn/view/pages/index/jiansuo.html）",
            "中国盐业协会（https://www.cnsalt.cn）",
            "国家统计局网站（https://www.stats.gov.cn）",
            "国家能源局网站（https://www.nea.gov.cn）",
            "国家市场监督管理总局（https://www.samr.gov.cn）",
            "国家农业农村部网站（https://www.moa.gov.cn）",
            "中华人民共和国住房和城乡建设部（https://www.mohurd.gov.cn）",
            "全国建筑市场监管公共服务平台（https://jzsc.mohurd.gov.cn）",
            "中华人民共和国人力资源和社会保障部网站（https://www.mohrss.gov.cn/hsearch/）",
            "中国电力企业联合会网站（https://cec.org.cn）",
            "国家药品监督管理局（https://www.nmpa.gov.cn）"
        ]

        # 定义需要人工核查的网页
        manual_chapters = [
            "国家企业信用信息公示系统（http://www.gsxt.gov.cn/index.html）",
            "国家税务总局重大税收违法案件信息公布栏（https://www.chinatax.gov.cn）",
            "信用中国网（http://www.creditchina.gov.cn）",
            "中国裁判文书网（http://wenshu.court.gov.cn）",
            "中华人民共和国国家发展和改革委员会网站（http://www.ndrc.gov.cn）",
            "中国执行信息公开网（http://zxgk.court.gov.cn/shixin/）",
            "中华人民共和国应急管理部网站（http://www.mem.gov.cn）",
            "中华人民共和国工业和信息化部网站（http://www.miit.gov.cn）",
            "中国商务信用平台（http://www.bcpcn.com）",
            "全国行业信用公共服务平台（http://www.bcp12312.org.cn）",
            "中国证券监督管理委员会网站（http://www.csrc.gov.cn）",
            "证券期货市场失信记录查询平台（http://neris.csrc.gov.cn/shixinchaxun/）",
            "中华人民共和国自然资源部（http://www.mnr.gov.cn/）",
            "国家财政部网站（http://www.mof.gov.cn/index.htm）",
            "中国海关企业进出口信用信息公示平台（http://credit.customs.gov.cn）",
            "全国资源公共交易平台（http://www.ggzy.gov.cn）",
            "中华人民共和国海关总署（http://www.customs.gov.cn/）",
            "中华人民共和国交通运输部网站（https://www.mot.gov.cn/）",
            "政府采购严重违法失信行为信息记录（http://www.ccgp.gov.cn/cr/list）",
            "信用能源（https://xyny.nea.gov.cn）",
            "被执行人信息查询（https://zxgk.court.gov.cn/zhzxgk/）"
        ]

        # 添加自动化核查的网页
        for chapter_title in automated_chapters:
            # 添加章节标题
            self.document_engine.add_title(chapter_title, level=2)

            # 查找是否有对应的 URL
            found = False
            for url, screenshot_path in entity_results.items():
                # 获取模板信息
                template_info = self._get_template_info(url)
                if template_info.get("InsertContext") == chapter_title:
                    # 添加截图
                    if screenshot_path and Path(screenshot_path).exists():
                        # 添加水印
                        watermarked_path = self._add_watermark(screenshot_path, entity)
                        self.document_engine.add_image(Path(watermarked_path))
                    found = True
                    break

        # 添加人工核查的网页
        for chapter_title in manual_chapters:
            # 添加章节标题
            self.document_engine.add_title(chapter_title, level=2)


        # 保存文档
        filename = f"{clean_filename(entity)} 诚信核查.docx"
        report_path = output_dir / filename
        self.document_engine.save_document(report_path)

        return report_path

    def _get_template_info(self, url: str) -> Dict:
        """根据URL获取模板信息

        Args:
            url: URL地址

        Returns:
            dict: 模板信息
        """
        if not self.templates:
            return {"name": "unknown", "description": "未知模板", "InsertContext": "网页核查"}

        for template_name, template in self.templates.items():
            if template.url_pattern and "{}" in template.url_pattern:
                # 检查URL是否匹配模板模式（简化匹配：只检查是否包含模板模式的固定部分）
                fixed_part = template.url_pattern.split("{}")[0]
                if fixed_part in url:
                    return {
                        "name": template_name,
                        "description": template.description,
                        "InsertContext": template.InsertContext if hasattr(template, "InsertContext") else "网页核查"
                    }

        return {"name": "unknown", "description": "未知模板", "InsertContext": "网页核查"}

    def _add_watermark(self, image_path: str, entity: str) -> str:
        """为图片添加水印

        Args:
            image_path: 原始图片路径
            entity: 主体名称

        Returns:
            str: 带水印的图片路径
        """
        from datetime import datetime
        current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        watermark_text = f"{entity} - {current_time}"

        # 生成输出路径
        from pathlib import Path
        img_path = Path(image_path)
        output_path = img_path.parent / f"{img_path.stem}_watermarked{img_path.suffix}"

        # 添加水印
        self.image_engine.add_watermark(
            img_path,
            watermark_text,
            output_path
        )

        return str(output_path)

    def generate_search_engine_report(
        self,
        search_results: Dict[str, Dict[str, str]],
        output_path: Optional[Path] = None
    ) -> Dict[str, Path]:
        """生成搜索引擎查询结果的报告

        Args:
            search_results: 搜索引擎查询结果 {主体名称: {关键词: 截图路径}}
            output_path: 报告输出路径

        Returns:
            dict: {主体名称: 报告文件路径}
        """
        output_dir = output_path or self.config.reports_dir
        output_dir.mkdir(parents=True, exist_ok=True)

        report_paths = {}

        for entity, entity_results in search_results.items():
            try:
                report_path = self._generate_single_search_engine_report(entity, entity_results, output_dir)
                report_paths[entity] = report_path
                logger.info(f"主体 '{entity}' 搜索引擎查询报告生成成功: {report_path}")
            except Exception as e:
                logger.error(f"主体 '{entity}' 搜索引擎查询报告生成失败: {e}")

        return report_paths

    def _generate_single_search_engine_report(
        self,
        entity: str,
        entity_results: Dict[str, str],
        output_dir: Path
    ) -> Path:
        """为单个主体生成搜索引擎查询结果报告

        Args:
            entity: 主体名称
            entity_results: 主体的搜索结果 {关键词: 截图路径}
            output_dir: 输出目录

        Returns:
            Path: 报告文件路径
        """
        from everify.utils.file import clean_filename

        # 创建文档
        from datetime import datetime
        current_date = datetime.now().strftime("%Y%m%d")
        title = f"{entity}搜索引擎查询报告{current_date}"
        self.document_engine.create_document(title)

        # 添加标题
        self.document_engine.add_title(title, level=1)

        # 定义需要搜索的关键词
        search_keywords = ['舆情', '查封', '冻结', '收购']

        # 添加搜索结果章节
        for keyword in search_keywords:
            # 添加章节标题
            chapter_title = f"公司{keyword}情况"
            self.document_engine.add_title(chapter_title, level=2)

            # 添加搜索结果截图
            if keyword in entity_results:
                screenshot_path = entity_results[keyword]
                if screenshot_path and Path(screenshot_path).exists():
                    # 添加水印
                    watermarked_path = self._add_watermark(screenshot_path, entity)
                    self.document_engine.add_image(Path(watermarked_path))
                else:
                    self.document_engine.add_paragraph("未找到截图")
            else:
                self.document_engine.add_paragraph("未找到搜索结果")

        # 保存文档
        filename = f"{clean_filename(entity)} 搜索引擎查询报告.docx"
        report_path = output_dir / filename
        self.document_engine.save_document(report_path)

        return report_path

    def insert_manual_screenshots(self, entity: str, report_path: Path, screenshot_dir: Path, all_entities: List[str]):
        """插入人工核查的截图

        Args:
            entity: 主体名称
            report_path: 报告文件路径
            screenshot_dir: 截图文件夹路径（所有截图都在这个目录下）
            all_entities: 所有需要核查的主体列表（保持原始输入顺序）
        """
        # 获取所有截图文件
        screenshots = list(screenshot_dir.glob("*.png")) + list(screenshot_dir.glob("*.jpg")) + list(screenshot_dir.glob("*.jpeg")) + list(screenshot_dir.glob("*.bmp"))
        if not screenshots:
            logger.warning(f"未找到截图文件: {screenshot_dir}")
            return

        # 按文件名排序（按时间戳顺序，确保顺序正确）
        screenshots.sort(key=lambda x: x.name)

        # 定义需要人工核查的网页（与 report_generator.py 中的一致）
        manual_chapters = [
            "国家企业信用信息公示系统（http://www.gsxt.gov.cn/index.html）",
            "国家税务总局重大税收违法案件信息公布栏（https://www.chinatax.gov.cn）",
            "信用中国网（http://www.creditchina.gov.cn）",
            "中国裁判文书网（http://wenshu.court.gov.cn）",
            "中华人民共和国国家发展和改革委员会网站（http://www.ndrc.gov.cn）",
            "中国执行信息公开网（http://zxgk.court.gov.cn/shixin/）",
            "中华人民共和国应急管理部网站（http://www.mem.gov.cn）",
            "中华人民共和国工业和信息化部网站（http://www.miit.gov.cn）",
            "中国商务信用平台（http://www.bcpcn.com）",
            "全国行业信用公共服务平台（http://www.bcp12312.org.cn）",
            "中国证券监督管理委员会网站（http://www.csrc.gov.cn）",
            "证券期货市场失信记录查询平台（http://neris.csrc.gov.cn/shixinchaxun/）",
            "中华人民共和国自然资源部（http://www.mnr.gov.cn/）",
            "国家财政部网站（http://www.mof.gov.cn/index.htm）",
            "中国海关企业进出口信用信息公示平台（http://credit.customs.gov.cn）",
            "全国资源公共交易平台（http://www.ggzy.gov.cn）",
            "中华人民共和国海关总署（http://www.customs.gov.cn/）",
            "中华人民共和国交通运输部网站（https://www.mot.gov.cn/）",
            "政府采购严重违法失信行为信息记录（http://www.ccgp.gov.cn/cr/list）",
            "信用能源（https://xyny.nea.gov.cn）",
            "被执行人信息查询（https://zxgk.court.gov.cn/zhzxgk/）"
        ]

        # 确定当前主体在列表中的索引
        try:
            entity_index = all_entities.index(entity)
        except ValueError:
            logger.warning(f"主体 '{entity}' 不在主体列表中")
            return

        # 计算每个主体每个网页的截图索引
        # 截图顺序：网页1主体1 → 网页1主体2 → 网页1主体3 → 网页1主体4 → 网页1主体5 → 网页2主体1 → ...
        entity_screenshots = []
        for web_index in range(len(manual_chapters)):
            # 计算该网页该主体的截图索引
            screenshot_index = web_index * len(all_entities) + entity_index
            if screenshot_index < len(screenshots):
                entity_screenshots.append(screenshots[screenshot_index])
            else:
                entity_screenshots.append(None)

        # 打开报告文件并插入截图
        from docx import Document
        from docx.shared import Inches

        try:
            doc = Document(report_path)

            # 遍历每个章节并插入对应的截图
            for i, chapter_title in enumerate(manual_chapters):
                # 查找章节标题所在的段落
                found = False
                for para in doc.paragraphs:
                    if para.style.name.startswith("Heading 2") and chapter_title in para.text:
                        # 插入截图（如果有对应的截图）
                        if i < len(entity_screenshots) and entity_screenshots[i] is not None:
                            para.add_run().add_picture(str(entity_screenshots[i]), width=Inches(6))
                        found = True
                        break

                if not found:
                    logger.warning(f"未找到章节标题: {chapter_title}")

            # 保存修改后的报告
            doc.save(report_path)
            logger.info(f"主体 '{entity}' 人工核查截图插入完成: {report_path}")

        except Exception as e:
            logger.error(f"插入人工核查截图失败: {e}")