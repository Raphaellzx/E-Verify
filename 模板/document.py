from docx import Document
from docx.shared import Inches
import os
from tqdm import tqdm

def insert_screenshots_to_word(doc_paths, screenshot_folder, output_folder,
                             file_count, items_per_file, pics_per_item):
    """
    向Word文档中插入截图

    参数:
        doc_paths: 模板文档路径列表
        screenshot_folder: 截图文件夹路径
        output_folder: 输出文件夹路径
        file_count: 输出文件数量
        items_per_file: 每个文件中的核查项目数量
        pics_per_item: 每个核查项目插入的图片数量
    """
    # 确保输出文件夹存在
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # 获取所有截图并排序（按文件名数字排序）
    screenshots = [f for f in os.listdir(screenshot_folder) if f.endswith(('.png', '.jpg', '.jpeg', '.bmp'))]
    # 按文件名中的数字排序，确保截图顺序正确
    screenshots.sort(key=lambda x: int(''.join(filter(str.isdigit, x))) if any(c.isdigit() for c in x) else 0)
    total_needed = file_count * items_per_file * pics_per_item
    if len(screenshots) < total_needed:
        print(f"警告：截图数量不足，需要{total_needed}张，实际只有{len(screenshots)}张")

    # 计算总任务数
    total_tasks = file_count * items_per_file * pics_per_item

    # 使用 tqdm 创建进度条，减少输出
    with tqdm(total=total_tasks, desc="正在插入截图", unit="张", ncols=80, leave=True, mininterval=1.0) as pbar:
        # 遍历每个输出文件
        for file_idx in range(file_count):
            # 获取当前文件的模板
            if file_idx >= len(doc_paths):
                print(f"警告：模板文档不足，跳过文件{file_idx+1}")
                continue

            template_path = doc_paths[file_idx]

            # 检查文件是否存在
            if not os.path.exists(template_path):
                continue

            # 检查文件是否为有效的docx文件
            try:
                doc = Document(template_path)
            except Exception as e:
                continue

            # 遍历当前文件中的每个核查项目
            for item_idx in range(items_per_file):
                # 计算当前项目的起始截图索引 - 按照项目→文件→图片的顺序
                start_idx = (item_idx * file_count + file_idx) * pics_per_item
                if start_idx >= len(screenshots):
                    # 跳过该项目的所有图片任务
                    pbar.update(pics_per_item)
                    continue

                # 查找当前项目对应的段落（每个项目占1个段落）
                if item_idx >= len(doc.paragraphs):
                    pbar.update(pics_per_item)
                    continue
                para = doc.paragraphs[item_idx]

                # 为当前项目插入指定数量的图片
                for pic_idx in range(pics_per_item):
                    screenshot_idx = start_idx + pic_idx
                    if screenshot_idx >= len(screenshots):
                        pbar.update(1)
                        continue

                    screenshot_path = os.path.join(screenshot_folder, screenshots[screenshot_idx])
                    # 插入换行和图片
                    para.add_run('\n')
                    run = para.add_run()
                    run.add_picture(screenshot_path, width=Inches(6))
                    pbar.update(1)

            # 保存处理后的文档
            output_doc_path = os.path.join(output_folder, f"处理后_文件{file_idx+1}.docx")
            doc.save(output_doc_path)

if __name__ == "__main__":
    '''
    ***运行前，请先安装依赖：pip install python-docx

    ***截图命名建议：包含数字序号（如screenshot_1.png, pic2.jpg等），程序会按数字排序
    ***截图顺序规则：按项目→文件→图片的层级排序
      例如：项目1文件1图片1、项目1文件1图片2、项目1文件2图片1、项目1文件2图片2、项目2文件1图片1...

    ***请将模板文档与本程序放在合适路径下
    '''
    
    # -------------------------- 可自主设定的参数 --------------------------
    file_count = 6          # 输出文件数量
    items_per_file = 36     # 每个文件中的核查项目数量
    pics_per_item = 1     # 每个核查项目插入的图片数量
    
    # 模板文档路径列表（可指定不同模板，长度不小于file_count）
    import os
    current_dir = os.path.dirname(os.path.abspath(__file__))
    template_filename = 'MODEL.docx'
    template_path = os.path.join(current_dir, template_filename)
    print(f"找到模板文件: {template_path}")
    doc_paths = [template_path for _ in range(file_count)]  # 使用相同模板

    # 截图文件夹路径
    screenshot_folder = r'D:\HONOR\HONOR Share\Screenshot'

    # 处理后文档的输出文件夹路径
    output_folder = r"D:\Coding\Python\Picture"
    # ----------------------------------------------------------------------
    
    # 执行插入操作
    insert_screenshots_to_word(
        doc_paths=doc_paths,
        screenshot_folder=screenshot_folder,
        output_folder=output_folder,
        file_count=file_count,
        items_per_file=items_per_file,
        pics_per_item=pics_per_item
    )